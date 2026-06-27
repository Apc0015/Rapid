"""
infrastructure/document_engine.py — Document Generation Engine.

Produces structured .docx reports from project data and stores them
in the project's document library.

Supported report types
──────────────────────
  status_report  — Executive project status (health, KPIs, milestones, risks)
  risk_report    — Full risk assessment with impact matrix
  pipeline_report — Sales pipeline snapshot with deal table
  kpi_report     — KPI scorecard with trend indicators
  custom         — Caller-assembled sections, no fixed template

Output
──────
  - .docx file written to data/documents/projects/{tenant_id}/{project_id}/
  - Row inserted/updated in project_documents table
  - GraphNode (COMMUNICATION) created via NodeIngestionPipeline

Usage
─────
    from infrastructure.document_engine import DocumentGenerator

    gen = DocumentGenerator(project_context)
    result = await gen.generate("status_report")
    # → DocumentResult(doc_id="...", file_path="...", title="...", pages=4)
"""

from __future__ import annotations

import logging
import os
import sqlite3
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from infrastructure.project_context import ProjectContext

logger = logging.getLogger(__name__)


# ── Result dataclass ──────────────────────────────────────────────────────────

@dataclass
class DocumentResult:
    doc_id:     str
    title:      str
    file_path:  str
    file_format: str    = "docx"
    report_type: str    = "status_report"
    pages:       int    = 0
    produced_by: str    = "document_engine"
    created_at:  str    = field(default_factory=lambda: datetime.utcnow().isoformat())
    error:       Optional[str] = None

    @property
    def success(self) -> bool:
        return self.error is None and bool(self.file_path) and Path(self.file_path).exists()


# ── Section dataclass (for custom reports) ────────────────────────────────────

@dataclass
class ReportSection:
    """A named section with optional table data."""
    heading:   str
    body:      str                        = ""
    table:     Optional[list[list[str]]] = None   # rows of cells
    level:     int                        = 1      # heading level (1=H1, 2=H2)


# ── DocumentGenerator ─────────────────────────────────────────────────────────

class DocumentGenerator:
    """
    Generates structured .docx reports from project data.

    Each public method corresponds to a report type and returns a DocumentResult.
    All generated files are stored in the project document folder and registered
    in the project_documents table.
    """

    _DOCS_BASE = Path("data/documents/projects")

    def __init__(self, project_context: ProjectContext):
        self.ctx        = project_context
        self.project_id = project_context.project_id
        self.tenant_id  = project_context.tenant_id
        self._db_path   = self._resolve_db_path()
        self._doc_dir   = (
            self._DOCS_BASE / self.tenant_id / self.project_id
        )
        self._doc_dir.mkdir(parents=True, exist_ok=True)

    # ── Public dispatch ───────────────────────────────────────────────────────

    async def generate(
        self,
        report_type:  str = "status_report",
        title:        Optional[str] = None,
        produced_by:  str = "document_engine",
        custom_sections: Optional[list[ReportSection]] = None,
    ) -> DocumentResult:
        """
        Generate a report of the given type.

        report_type: 'status_report' | 'risk_report' | 'pipeline_report' |
                     'kpi_report'    | 'custom'
        """
        try:
            if report_type == "status_report":
                result = await self._build_status_report(title, produced_by)
            elif report_type == "risk_report":
                result = await self._build_risk_report(title, produced_by)
            elif report_type == "pipeline_report":
                result = await self._build_pipeline_report(title, produced_by)
            elif report_type == "kpi_report":
                result = await self._build_kpi_report(title, produced_by)
            elif report_type == "custom" and custom_sections:
                result = await self._build_custom_report(
                    title or "Custom Report", custom_sections, produced_by
                )
            else:
                raise ValueError(f"Unknown report_type '{report_type}'")

            # Register in project_documents table + trigger graph node
            self._register_document(result)
            return result

        except Exception as e:
            logger.error(f"[DocumentEngine] Generation failed ({report_type}): {e}")
            return DocumentResult(
                doc_id    = str(uuid.uuid4()),
                title     = title or report_type,
                file_path = "",
                error     = str(e),
            )

    # ── Report builders ───────────────────────────────────────────────────────

    async def _build_status_report(
        self,
        title:       Optional[str],
        produced_by: str,
    ) -> DocumentResult:
        """Executive project status report."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        ctx   = self.ctx
        meta  = ctx.metadata or {}
        now   = datetime.utcnow().strftime("%B %d, %Y")
        title = title or f"{ctx.project_name} — Status Report"

        doc = Document()
        self._set_document_style(doc)

        # ── Cover / title
        self._add_title(doc, title)
        self._add_subtitle(doc, f"Project: {ctx.project_name} | Dept: {ctx.dept_id.upper()} | Generated: {now}")

        # ── Executive Summary
        self._add_heading(doc, "Executive Summary", level=1)
        health   = meta.get("health_status", "unknown").upper()
        comp_pct = meta.get("completion_pct", 0)
        bt       = meta.get("budget_total", 0) or 0
        bs       = meta.get("budget_spent", 0) or 0
        util_pct = (bs / bt * 100) if bt else 0
        summary  = (
            f"Project health is {health}. Overall completion stands at {comp_pct:.0f}%. "
            f"Budget utilization is {util_pct:.1f}% "
            f"({self._fmt_currency(bs)} spent of {self._fmt_currency(bt)}). "
            f"Target end date: {meta.get('target_end_date', 'TBD')}."
        )
        doc.add_paragraph(summary)

        # ── KPI Scorecard
        self._add_heading(doc, "KPI Scorecard", level=1)
        kpis = ctx.kpi_summary or []
        if kpis:
            rows = [["KPI", "Current", "Target", "Unit", "Status", "Trend"]]
            for k in kpis:
                rows.append([
                    str(k.get("kpi_name", "")),
                    str(k.get("current_value", "")),
                    str(k.get("target_value", "")),
                    str(k.get("unit", "")),
                    str(k.get("status", "")),
                    str(k.get("trend", "—")),
                ])
            self._add_table(doc, rows)
        else:
            doc.add_paragraph("No KPI data available.")

        # ── Milestones
        self._add_heading(doc, "Milestone Progress", level=1)
        milestones = self._fetch_milestones()
        if milestones:
            rows = [["Milestone", "Due Date", "Status", "Priority"]]
            for m in milestones:
                rows.append([m["name"], m["due_date"] or "TBD", m["status"], m["priority"] or "—"])
            self._add_table(doc, rows)
        else:
            doc.add_paragraph("No milestones recorded.")

        # ── Open Risks
        self._add_heading(doc, "Risk Summary", level=1)
        risks = self._fetch_open_risks()
        if risks:
            rows = [["Risk", "Probability", "Impact", "Status"]]
            for r in risks:
                rows.append([r["title"], r["probability"] or "—", r["impact"] or "—", r["status"]])
            self._add_table(doc, rows)
        else:
            doc.add_paragraph("No open risks.")

        # ── Footer note
        doc.add_paragraph()
        p = doc.add_paragraph(f"Generated by RAPID Document Engine · {now}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        file_path = self._save(doc, "status_report")
        return DocumentResult(
            doc_id      = str(uuid.uuid4()),
            title       = title,
            file_path   = str(file_path),
            report_type = "status_report",
            produced_by = produced_by,
            pages       = max(1, len(milestones) // 8 + len(kpis) // 10 + 2),
        )

    async def _build_risk_report(
        self,
        title:       Optional[str],
        produced_by: str,
    ) -> DocumentResult:
        """Full risk assessment report."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        now   = datetime.utcnow().strftime("%B %d, %Y")
        title = title or f"{self.ctx.project_name} — Risk Assessment"

        doc = Document()
        self._set_document_style(doc)
        self._add_title(doc, title)
        self._add_subtitle(doc, f"Project: {self.ctx.project_name} | Generated: {now}")

        # All risks from DB
        risks = self._fetch_all_risks()

        # ── Risk Overview
        self._add_heading(doc, "Risk Overview", level=1)
        open_risks = [r for r in risks if r["status"] == "open"]
        high_impact = [r for r in open_risks if r.get("impact") == "high"]
        doc.add_paragraph(
            f"Total risks tracked: {len(risks)}. Open: {len(open_risks)}. "
            f"High-impact open risks: {len(high_impact)}."
        )

        # ── Risk Matrix table
        self._add_heading(doc, "Risk Register", level=1)
        if risks:
            rows = [["Risk Title", "Category", "Probability", "Impact", "Score", "Status", "Mitigation"]]
            for r in risks:
                rows.append([
                    r["title"],
                    r.get("category") or "—",
                    r.get("probability") or "—",
                    r.get("impact") or "—",
                    str(r.get("risk_score") or "—"),
                    r["status"],
                    (r.get("mitigation_plan") or "None")[:80],
                ])
            self._add_table(doc, rows)
        else:
            doc.add_paragraph("No risks recorded in this project.")

        # ── High-impact risks deep dive
        if high_impact:
            self._add_heading(doc, "High-Impact Risk Details", level=1)
            for r in high_impact:
                self._add_heading(doc, r["title"], level=2)
                details = (
                    f"Category: {r.get('category', 'Unknown')} | "
                    f"Probability: {r.get('probability', '—')} | "
                    f"Impact: {r.get('impact', '—')} | "
                    f"Score: {r.get('risk_score', '—')}"
                )
                doc.add_paragraph(details)
                if r.get("mitigation_plan"):
                    doc.add_paragraph(f"Mitigation: {r['mitigation_plan']}")

        p = doc.add_paragraph(f"Generated by RAPID Document Engine · {now}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        file_path = self._save(doc, "risk_report")
        return DocumentResult(
            doc_id      = str(uuid.uuid4()),
            title       = title,
            file_path   = str(file_path),
            report_type = "risk_report",
            produced_by = produced_by,
            pages       = max(1, len(risks) // 6 + 2),
        )

    async def _build_pipeline_report(
        self,
        title:       Optional[str],
        produced_by: str,
    ) -> DocumentResult:
        """Sales pipeline snapshot."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        now   = datetime.utcnow().strftime("%B %d, %Y")
        title = title or f"{self.ctx.project_name} — Pipeline Report"

        doc = Document()
        self._set_document_style(doc)
        self._add_title(doc, title)
        self._add_subtitle(doc, f"Project: {self.ctx.project_name} | Generated: {now}")

        deals = self._fetch_pipeline()

        # ── Pipeline Summary
        self._add_heading(doc, "Pipeline Summary", level=1)
        total_value    = sum(float(d.get("value") or 0) for d in deals)
        avg_prob       = (sum(float(d.get("probability") or 0) for d in deals) / len(deals)) if deals else 0
        weighted_value = sum(
            float(d.get("value") or 0) * float(d.get("probability") or 0) / 100
            for d in deals
        )
        doc.add_paragraph(
            f"Total deals: {len(deals)} | Total pipeline value: {self._fmt_currency(total_value)} | "
            f"Average probability: {avg_prob:.0f}% | "
            f"Weighted (expected) value: {self._fmt_currency(weighted_value)}"
        )

        # ── Stage breakdown
        if deals:
            stage_map: dict[str, list] = {}
            for d in deals:
                stage = d.get("stage") or "Unknown"
                stage_map.setdefault(stage, []).append(d)

            self._add_heading(doc, "Pipeline by Stage", level=2)
            stage_rows = [["Stage", "Deals", "Total Value", "Avg Probability"]]
            for stage, stage_deals in sorted(stage_map.items()):
                sv = sum(float(d.get("value") or 0) for d in stage_deals)
                sp = sum(float(d.get("probability") or 0) for d in stage_deals) / len(stage_deals)
                stage_rows.append([stage, str(len(stage_deals)), self._fmt_currency(sv), f"{sp:.0f}%"])
            self._add_table(doc, stage_rows)

        # ── Full deal table
        self._add_heading(doc, "Deal Register", level=1)
        if deals:
            rows = [["Customer", "Stage", "Value", "Probability", "Close Date", "Owner"]]
            for d in sorted(deals, key=lambda x: float(x.get("value") or 0), reverse=True):
                rows.append([
                    str(d.get("customer_name") or "—"),
                    str(d.get("stage") or "—"),
                    self._fmt_currency(d.get("value") or 0),
                    f"{d.get('probability') or 0}%",
                    str(d.get("close_date") or "TBD"),
                    str(d.get("owner") or "—"),
                ])
            self._add_table(doc, rows)
        else:
            doc.add_paragraph("No pipeline deals recorded.")

        p = doc.add_paragraph(f"Generated by RAPID Document Engine · {now}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        file_path = self._save(doc, "pipeline_report")
        return DocumentResult(
            doc_id      = str(uuid.uuid4()),
            title       = title,
            file_path   = str(file_path),
            report_type = "pipeline_report",
            produced_by = produced_by,
            pages       = max(1, len(deals) // 10 + 2),
        )

    async def _build_kpi_report(
        self,
        title:       Optional[str],
        produced_by: str,
    ) -> DocumentResult:
        """KPI scorecard report with all KPIs and trend indicators."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        now   = datetime.utcnow().strftime("%B %d, %Y")
        title = title or f"{self.ctx.project_name} — KPI Scorecard"

        doc = Document()
        self._set_document_style(doc)
        self._add_title(doc, title)
        self._add_subtitle(doc, f"Project: {self.ctx.project_name} | Generated: {now}")

        kpis = self._fetch_all_kpis()

        # ── KPI Status Overview
        self._add_heading(doc, "KPI Status Overview", level=1)
        on_track  = sum(1 for k in kpis if k.get("status") == "on_track")
        at_risk   = sum(1 for k in kpis if k.get("status") == "at_risk")
        off_track = sum(1 for k in kpis if k.get("status") == "off_track")
        doc.add_paragraph(
            f"Total KPIs tracked: {len(kpis)} | "
            f"On track: {on_track} | At risk: {at_risk} | Off track: {off_track}"
        )

        # ── KPI table (all)
        self._add_heading(doc, "KPI Details", level=1)
        if kpis:
            rows = [["KPI", "Current", "Target", "Unit", "Status", "Trend", "Period"]]
            for k in sorted(kpis, key=lambda x: x.get("status", "")):
                trend_icon = {"improving": "↑", "declining": "↓", "stable": "→"}.get(
                    k.get("trend", ""), "—"
                )
                rows.append([
                    str(k.get("kpi_name") or "—"),
                    str(k.get("current_value") or "—"),
                    str(k.get("target_value") or "—"),
                    str(k.get("unit") or "—"),
                    str(k.get("status") or "—"),
                    trend_icon,
                    str(k.get("period") or "—"),
                ])
            self._add_table(doc, rows)
        else:
            doc.add_paragraph("No KPI data recorded.")

        p = doc.add_paragraph(f"Generated by RAPID Document Engine · {now}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        file_path = self._save(doc, "kpi_report")
        return DocumentResult(
            doc_id      = str(uuid.uuid4()),
            title       = title,
            file_path   = str(file_path),
            report_type = "kpi_report",
            produced_by = produced_by,
            pages       = max(1, len(kpis) // 12 + 1),
        )

    async def _build_custom_report(
        self,
        title:           str,
        sections:        list[ReportSection],
        produced_by:     str,
    ) -> DocumentResult:
        """Assemble a report from caller-supplied ReportSection objects."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        now = datetime.utcnow().strftime("%B %d, %Y")
        doc = Document()
        self._set_document_style(doc)
        self._add_title(doc, title)
        self._add_subtitle(doc, f"Project: {self.ctx.project_name} | Generated: {now}")

        for section in sections:
            self._add_heading(doc, section.heading, level=section.level)
            if section.body:
                doc.add_paragraph(section.body)
            if section.table:
                self._add_table(doc, section.table)

        p = doc.add_paragraph(f"Generated by RAPID Document Engine · {now}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        file_path = self._save(doc, "custom_report")
        return DocumentResult(
            doc_id      = str(uuid.uuid4()),
            title       = title,
            file_path   = str(file_path),
            report_type = "custom",
            produced_by = produced_by,
            pages       = max(1, len(sections) // 3 + 1),
        )

    # ── Document styling helpers ──────────────────────────────────────────────

    def _set_document_style(self, doc) -> None:
        """Apply consistent base font to the document."""
        from docx.shared import Pt
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

    def _add_title(self, doc, text: str) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    def _add_subtitle(self, doc, text: str) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_paragraph()

    def _add_heading(self, doc, text: str, level: int = 1) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        sizes = {1: 14, 2: 12, 3: 11}
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64) if level == 1 else RGBColor(0x26, 0x4F, 0x78)

    def _add_table(self, doc, rows: list[list[str]]) -> None:
        """Add a styled table. First row is the header."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        if not rows:
            return
        n_cols = len(rows[0])
        table  = doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_text)
                para = cell.paragraphs[0]
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                # Header row background
                if i == 0:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd   = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "1F3564")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:val"), "clear")
                    tc_pr.append(shd)

        doc.add_paragraph()

    # ── DB query helpers ──────────────────────────────────────────────────────

    def _connect_ro(self) -> sqlite3.Connection:
        conn = sqlite3.connect(f"file:{self._db_path}?mode=ro", uri=True, timeout=10)
        conn.row_factory = sqlite3.Row
        return conn

    def _fetch_milestones(self) -> list[dict]:
        try:
            conn = self._connect_ro()
            rows = conn.execute(
                "SELECT name, due_date, status, priority FROM project_milestones ORDER BY due_date"
            ).fetchall()
            conn.close()
            return [dict(r) for r in rows]
        except Exception:
            return []

    def _fetch_open_risks(self, limit: int = 20) -> list[dict]:
        try:
            conn = self._connect_ro()
            rows = conn.execute(
                "SELECT title, probability, impact, status FROM project_risks "
                "WHERE status='open' ORDER BY risk_score DESC LIMIT ?",
                (limit,),
            ).fetchall()
            conn.close()
            return [dict(r) for r in rows]
        except Exception:
            return []

    def _fetch_all_risks(self) -> list[dict]:
        try:
            conn = self._connect_ro()
            rows = conn.execute(
                "SELECT title, category, probability, impact, risk_score, status, mitigation_plan "
                "FROM project_risks ORDER BY risk_score DESC"
            ).fetchall()
            conn.close()
            return [dict(r) for r in rows]
        except Exception:
            return []

    def _fetch_pipeline(self) -> list[dict]:
        try:
            conn = self._connect_ro()
            rows = conn.execute(
                "SELECT customer_name, stage, value, probability, close_date, owner "
                "FROM project_pipeline ORDER BY value DESC"
            ).fetchall()
            conn.close()
            return [dict(r) for r in rows]
        except Exception:
            return []

    def _fetch_all_kpis(self) -> list[dict]:
        try:
            conn = self._connect_ro()
            rows = conn.execute(
                "SELECT kpi_name, current_value, target_value, unit, status, trend, period "
                "FROM project_kpis ORDER BY status"
            ).fetchall()
            conn.close()
            return [dict(r) for r in rows]
        except Exception:
            return []

    # ── File saving ───────────────────────────────────────────────────────────

    def _save(self, doc, report_type: str) -> Path:
        ts   = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        name = f"{report_type}_{ts}.docx"
        path = self._doc_dir / name
        doc.save(str(path))
        logger.info(f"[DocumentEngine] Saved {name} ({path.stat().st_size // 1024}KB)")
        return path

    # ── Project registration ──────────────────────────────────────────────────

    def _register_document(self, result: DocumentResult) -> None:
        """Insert the document into project_documents and create a graph node."""
        try:
            conn = sqlite3.connect(self._db_path, timeout=10)
            conn.execute(
                """
                INSERT OR REPLACE INTO project_documents
                    (doc_id, title, doc_type, skill_used, file_path,
                     file_format, produced_by, status, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    result.doc_id,
                    result.title,
                    result.report_type,
                    "document_engine",
                    result.file_path,
                    result.file_format,
                    result.produced_by,
                    "final",
                    result.created_at,
                ),
            )
            conn.commit()
            conn.close()
            logger.debug(f"[DocumentEngine] Registered doc {result.doc_id} in project DB")
        except Exception as e:
            logger.warning(f"[DocumentEngine] Could not register document: {e}")

        # Create knowledge graph node (COMMUNICATION type)
        try:
            from infrastructure.node_ingestion import NodeIngestionPipeline
            pipeline = NodeIngestionPipeline(self._db_path, self.project_id, self.tenant_id)
            pipeline.ingest_single_row(
                table_name = "project_documents",
                row_data   = {
                    "doc_id":      result.doc_id,
                    "title":       result.title,
                    "doc_type":    result.report_type,
                    "skill_used":  "document_engine",
                    "file_format": result.file_format,
                    "status":      "final",
                    "produced_by": result.produced_by,
                },
                source_id  = result.doc_id,
            )
        except Exception as e:
            logger.debug(f"[DocumentEngine] Graph node creation skipped: {e}")

    # ── Utility ───────────────────────────────────────────────────────────────

    def _resolve_db_path(self) -> str:
        """Locate the project's SQLite database."""
        from agents.base.project_aware_mixin import ProjectAwareMixin
        m = ProjectAwareMixin()
        path = m.get_project_db_path(self.ctx)
        return path or ""

    @staticmethod
    def _fmt_currency(val: Any) -> str:
        try:
            return f"${float(val):,.0f}"
        except (TypeError, ValueError):
            return str(val)


# ── Convenience factory ───────────────────────────────────────────────────────

def get_document_generator(project_context: ProjectContext) -> DocumentGenerator:
    """Create a DocumentGenerator for the given project context."""
    return DocumentGenerator(project_context)
