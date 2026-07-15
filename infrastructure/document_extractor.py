"""Safe text extraction for organization knowledge uploads with optional OCR."""
from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path


class DocumentExtractionError(ValueError):
    """Safe extraction error returned to an uploader."""


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    method: str
    pages: int = 1


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        text = content.decode("utf-8", errors="replace")
        return ExtractedDocument(text=text, method="text")
    if suffix == ".json":
        try:
            value = json.loads(content.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as error:
            raise DocumentExtractionError("The JSON document is invalid") from error
        return ExtractedDocument(text=json.dumps(value, indent=2, default=str), method="json")
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        try:
            from docx import Document
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                paragraphs.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            text = "\n\n".join(paragraphs)
        except Exception as error:
            raise DocumentExtractionError("The Word document could not be parsed") from error
        if not text.strip():
            raise DocumentExtractionError("The Word document contains no extractable text")
        return ExtractedDocument(text=text, method="docx")
    if suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        return _ocr_image(content)
    raise DocumentExtractionError("Supported files are PDF, DOCX, TXT, Markdown, CSV, JSON, PNG, JPG, and TIFF")


def _extract_pdf(content: bytes) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as error:
        raise DocumentExtractionError("The PDF could not be parsed") from error
    if text:
        return ExtractedDocument(text=text, method="pdf_text", pages=len(reader.pages))
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(content, dpi=200)
        text = "\n\n".join(pytesseract.image_to_string(image) for image in images).strip()
    except (ImportError, OSError) as error:
        raise DocumentExtractionError("This scanned PDF requires the optional OCR runtime") from error
    if not text:
        raise DocumentExtractionError("OCR found no readable text in the PDF")
    return ExtractedDocument(text=text, method="pdf_ocr", pages=len(images))


def _ocr_image(content: bytes) -> ExtractedDocument:
    try:
        from PIL import Image
        import pytesseract
        text = pytesseract.image_to_string(Image.open(io.BytesIO(content))).strip()
    except (ImportError, OSError) as error:
        raise DocumentExtractionError("Image extraction requires the optional OCR runtime") from error
    if not text:
        raise DocumentExtractionError("OCR found no readable text in the image")
    return ExtractedDocument(text=text, method="image_ocr")
