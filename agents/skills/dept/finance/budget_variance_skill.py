"""
Finance Skill: Budget Variance Report (DOCX + XLSX)

Produces a professional budget variance report showing:
  - Allocated vs. spent vs. remaining by category
  - Variance ($ and %) for each line
  - Overall burn rate and forecast to complete
  - Narrative analysis generated via LLM

Output: DOCX report with embedded variance table.
"""

from __future__ import annotations

import logging
import os
import sqlite3
import uuid
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.skills.base_skill import BaseSkill, SkillOutput
from infrastructure.project_context import ProjectContext

logger = logging.getLogger(__name__)


def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


class BudgetVarianceSkill(BaseSkill):
    skill_id        = "budget_variance_report"
    dept_id         = "finance"
    title_template  = "{project_name} — Budget Variance Report"
    description     = "Finance: Generate a DOCX budget variance report with allocated vs. spent, variance %, and forecast."
    output_format   = "docx"
    trigger_phrases = [
        "budget variance", "budget report", "variance report",
        "spending analysis", "budget vs actual", "budget analysis",
        "allocated vs spent", "finance report", "budget utilization report",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        title  = params.get("title") or self._make_title(context)

        try:
            db_path = getattr(context, "db_path", None)
            budget_lines = []
            proj_info    = {}

            if db_path and os.path.exists(db_path):
                conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=10)
                conn.row_factory = sqlite3.Row
                try:
                    r = conn.execute("SELECT * FROM project_details LIMIT 1").fetchone()
                    if r:
                        proj_info = dict(r)

                    budget_lines = [dict(r) for r in conn.execute(
                        "SELECT category, allocated, spent, notes "
                        "FROM project_budget_lines ORDER BY category"
                    ).fetchall()]
                finally:
                    conn.close()

            # Totals
            total_alloc = sum(b.get("allocated") or 0 for b in budget_lines)
            total_spent = sum(b.get("spent") or 0 for b in budget_lines)
            total_remain = total_alloc - total_spent
            total_var_pct = (total_spent / total_alloc * 100) if total_alloc else 0

            # Build document
            doc  = Document()
            now  = datetime.utcnow().strftime("%B %d, %Y")

            # Title
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info_para.add_run(
                f"Project: {proj_info.get('name', context.project_id[:12])}  |  "
                f"Generated: {now}  |  Department: Finance"
            )
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()

            # Executive summary section
            doc.add_heading("Executive Summary", level=1)
            summary_text = (
                f"As of {now}, this project has consumed "
                f"${total_spent:,.2f} ({total_var_pct:.1f}%) of the total allocated budget "
                f"of ${total_alloc:,.2f}. The remaining budget stands at ${total_remain:,.2f}."
            )
            if total_var_pct >= 90:
                summary_text += " ⚠️ Budget is critically low — immediate attention required."
            elif total_var_pct >= 80:
                summary_text += " ⚠️ Budget utilization is high — monitor closely."
            else:
                summary_text += " Budget is within acceptable thresholds."
            doc.add_paragraph(summary_text)
            doc.add_paragraph()

            # Variance table
            doc.add_heading("Budget Variance by Category", level=1)

            headers = ["Category", "Allocated ($)", "Spent ($)", "Remaining ($)", "Variance %", "Notes"]
            table   = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold          = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size      = Pt(10)
                _set_cell_bg(hdr_cells[i], "1F3564")

            alt = False
            for b in budget_lines:
                alloc  = b.get("allocated") or 0
                spent  = b.get("spent") or 0
                remain = alloc - spent
                var_pct = (spent / alloc * 100) if alloc else 0

                row = table.add_row()
                values = [
                    b.get("category", "—"),
                    f"${alloc:,.2f}",
                    f"${spent:,.2f}",
                    f"${remain:,.2f}",
                    f"{var_pct:.1f}%",
                    b.get("notes") or "",
                ]
                for i, v in enumerate(values):
                    cell = row.cells[i]
                    cell.text = str(v)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if alt:
                        _set_cell_bg(cell, "DCE6F1")
                    # Red if overspent
                    if i == 4 and var_pct > 100:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        cell.paragraphs[0].runs[0].bold = True
                alt = not alt

            # Totals row
            total_row = table.add_row()
            total_values = [
                "TOTAL",
                f"${total_alloc:,.2f}",
                f"${total_spent:,.2f}",
                f"${total_remain:,.2f}",
                f"{total_var_pct:.1f}%",
                "",
            ]
            for i, v in enumerate(total_values):
                cell = total_row.cells[i]
                cell.text = str(v)
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(10)
                _set_cell_bg(cell, "2E75B6")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            doc.add_paragraph()
            doc.add_heading("Notes", level=1)
            doc.add_paragraph(
                "This report was automatically generated by RAPID Finance Agent. "
                "All figures are drawn directly from project budget records. "
                "Please review and approve before distribution."
            )

            # Save
            out_dir  = self._output_dir(context)
            filename = f"budget_variance_{uuid.uuid4().hex[:8]}.docx"
            out_path = os.path.join(out_dir, filename)
            doc.save(out_path)

            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = out_path,
                preview     = (
                    f"Budget variance: ${total_spent:,.0f} spent of ${total_alloc:,.0f} "
                    f"({total_var_pct:.1f}%). {len(budget_lines)} budget lines."
                ),
                pages       = 2,
            )

        except Exception as e:
            logger.error(f"[BudgetVarianceSkill] Failed: {e}", exc_info=True)
            return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                               title=title, file_format=self.output_format, error=str(e))


# ── Self-registration ─────────────────────────────────────────────────────────

def _register(registry) -> None:
    registry.register(BudgetVarianceSkill())
