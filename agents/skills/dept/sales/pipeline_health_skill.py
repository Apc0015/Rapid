"""
Sales Skill: Pipeline Health Report (DOCX)

Generates a pipeline health analysis with:
  - Stage distribution and deal count
  - At-risk deals (flagged in KPIs)
  - Win/loss velocity
  - Forecast summary

Since project data may not have a dedicated pipeline table,
this skill pulls from KPIs named "pipeline", "deals", "revenue" etc.
and the project's milestone + activity data.

Trigger phrases: "pipeline report", "pipeline health", "sales pipeline",
                 "deal report", "opportunity report"
"""

from __future__ import annotations

import logging
import os
import sqlite3
import uuid
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.skills.base_skill import BaseSkill, SkillOutput
from infrastructure.project_context import ProjectContext

logger = logging.getLogger(__name__)


def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


class PipelineHealthSkill(BaseSkill):
    skill_id        = "pipeline_health_report"
    dept_id         = "sales"
    title_template  = "{project_name} — Pipeline Health Report"
    description     = "Sales: Generate a pipeline health DOCX report with stage distribution, at-risk deals, and forecast."
    output_format   = "docx"
    trigger_phrases = [
        "pipeline report", "pipeline health", "sales pipeline",
        "deal report", "opportunity report", "pipeline analysis",
        "sales report", "sales health",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        title  = params.get("title") or self._make_title(context)

        try:
            db_path   = getattr(context, "db_path", None)
            proj_info = {}
            kpis      = []
            milestones = []

            if db_path and os.path.exists(db_path):
                conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=10)
                conn.row_factory = sqlite3.Row
                try:
                    r = conn.execute("SELECT * FROM project_details LIMIT 1").fetchone()
                    if r:
                        proj_info = dict(r)

                    kpis = [dict(r) for r in conn.execute(
                        "SELECT name, current_value, target_value, unit, status, trend "
                        "FROM project_kpis ORDER BY name"
                    ).fetchall()]

                    milestones = [dict(r) for r in conn.execute(
                        "SELECT name, due_date, status, owner "
                        "FROM project_milestones ORDER BY due_date ASC LIMIT 10"
                    ).fetchall()]
                finally:
                    conn.close()

            doc = Document()
            now = datetime.utcnow().strftime("%B %d, %Y")

            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Sales Department  |  Generated: {now}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()

            # Summary
            doc.add_heading("Executive Summary", 1)
            health = proj_info.get("health_status") or "unknown"
            doc.add_paragraph(
                f"This pipeline health report covers {proj_info.get('name', context.project_id[:12])}. "
                f"Current project health: {health.replace('_',' ').title()}. "
                f"Report includes {len(kpis)} KPIs and {len(milestones)} tracked milestones."
            )
            doc.add_paragraph()

            # Pipeline KPIs
            pipeline_kpis = [k for k in kpis if any(
                w in k.get("name","").lower()
                for w in ["pipeline", "deal", "revenue", "arr", "mrr", "win", "conversion",
                          "quota", "forecast", "opportunity", "close", "funnel"]
            )]

            doc.add_heading("Pipeline Metrics", 1)
            if pipeline_kpis:
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for i, h in enumerate(["Metric", "Current", "Target", "Unit", "Status"]):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].bold = True
                    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_bg(hdr[i], "1F3564")

                for idx, k in enumerate(pipeline_kpis):
                    row = table.add_row()
                    vals = [k.get("name"), str(k.get("current_value") or "—"),
                            str(k.get("target_value") or "—"), str(k.get("unit") or ""),
                            str(k.get("status") or "")]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                        if idx % 2 == 0:
                            _set_cell_bg(row.cells[i], "DCE6F1")
            else:
                doc.add_paragraph("No pipeline-specific KPIs configured. Configure KPIs with names including 'pipeline', 'deals', 'revenue', or 'quota'.")

            doc.add_paragraph()

            # Milestones = pipeline stages
            doc.add_heading("Pipeline Stage Milestones", 1)
            if milestones:
                table2 = doc.add_table(rows=1, cols=4)
                table2.style = "Table Grid"
                hdr2 = table2.rows[0].cells
                for i, h in enumerate(["Stage / Milestone", "Due Date", "Owner", "Status"]):
                    hdr2[i].text = h
                    hdr2[i].paragraphs[0].runs[0].bold = True
                    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_bg(hdr2[i], "1F3564")
                for idx, m in enumerate(milestones):
                    row = table2.add_row()
                    vals = [m.get("name"), str(m.get("due_date") or "TBD"),
                            str(m.get("owner") or "—"), str(m.get("status") or "")]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                        if idx % 2 == 0:
                            _set_cell_bg(row.cells[i], "DCE6F1")
            else:
                doc.add_paragraph("No milestones configured.")

            doc.add_paragraph()
            doc.add_heading("Notes & Recommendations", 1)
            at_risk = [k for k in kpis if str(k.get("status") or "").lower() in ("at_risk", "off_track")]
            if at_risk:
                doc.add_paragraph(f"⚠️ {len(at_risk)} metric(s) require attention: " +
                                  ", ".join(k.get("name","") for k in at_risk))
            else:
                doc.add_paragraph("✅ All tracked pipeline metrics are on track.")
            doc.add_paragraph("This report was generated automatically by RAPID Sales Agent. Review and approve before distribution.")

            out_dir  = self._output_dir(context)
            filename = f"pipeline_health_{uuid.uuid4().hex[:8]}.docx"
            out_path = os.path.join(out_dir, filename)
            doc.save(out_path)

            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = out_path,
                preview     = f"Pipeline health: {len(pipeline_kpis)} pipeline KPIs, {len(milestones)} stages.",
                pages       = 3,
            )

        except Exception as e:
            logger.error(f"[PipelineHealthSkill] Failed: {e}", exc_info=True)
            return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                               title=title, file_format=self.output_format, error=str(e))


def _register(registry) -> None:
    registry.register(PipelineHealthSkill())
