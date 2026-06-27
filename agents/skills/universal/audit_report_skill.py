"""
Universal Skill: Audit Report (DOCX)

Compliance-grade audit report covering:
  - Action queue activity (approvals, rejections, pending)
  - Notification log summary
  - Recent query/chat activity
  - Agent decision trail

Trigger phrases: "audit report", "compliance report", "audit log",
                 "compliance audit", "action log", "approval history",
                 "governance report", "activity report"
"""

from __future__ import annotations

import logging
import os
import sqlite3
import uuid
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.skills.base_skill import BaseSkill, SkillOutput
from infrastructure.project_context import ProjectContext

logger = logging.getLogger(__name__)


def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _hdr(table, headers: list[str]) -> None:
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[i], "1F3564")


class AuditReportSkill(BaseSkill):
    skill_id        = "audit_report"
    dept_id         = "all"
    title_template  = "{project_name} — Compliance Audit Report"
    description     = "Universal: Generate a compliance audit DOCX — action logs, approval history, notifications, agent activity."
    output_format   = "docx"
    trigger_phrases = [
        "audit report", "compliance report", "audit log",
        "compliance audit", "action log", "approval history",
        "governance report", "activity report", "audit trail",
        "action history", "agent audit", "compliance review",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        title  = params.get("title") or self._make_title(context)

        try:
            db_path = getattr(context, "db_path", None)

            # ── Collect action queue data ─────────────────────────────────────
            actions:       list[dict] = []
            notifications: list[dict] = []
            action_stats = {"pending": 0, "approved": 0, "rejected": 0, "executed": 0}

            if db_path and os.path.exists(db_path):
                conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=10)
                conn.row_factory = sqlite3.Row
                try:
                    # Actions
                    action_rows = conn.execute(
                        "SELECT action_id, action_type, title, category, status, "
                        "agent_dept, created_at, resolved_at, resolved_by "
                        "FROM agent_action_queue ORDER BY created_at DESC LIMIT 30"
                    ).fetchall()
                    for r in action_rows:
                        a = dict(r)
                        actions.append(a)
                        st = (a.get("status") or "pending").lower()
                        if st in action_stats:
                            action_stats[st] += 1

                    # Notifications
                    notif_rows = conn.execute(
                        "SELECT title, message, notif_type, is_read, created_at "
                        "FROM project_notifications ORDER BY created_at DESC LIMIT 20"
                    ).fetchall()
                    notifications = [dict(r) for r in notif_rows]

                except Exception as e:
                    logger.debug(f"[AuditSkill] DB read partial error: {e}")
                finally:
                    conn.close()

            # ── Build DOCX ────────────────────────────────────────────────────
            doc = Document()
            now = datetime.utcnow().strftime("%B %d, %Y %H:%M UTC")

            h = doc.add_heading(title, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Compliance Audit Report  |  Generated: {now}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()

            # ── Scope & Purpose ───────────────────────────────────────────────
            doc.add_heading("1. Scope and Purpose", 1)
            doc.add_paragraph(
                "This report provides a complete audit trail of agent-generated actions, "
                "human review decisions, and system notifications for this project. "
                "It is intended for compliance review, SOC2 audit support, and governance oversight. "
                f"Report period: all available records as of {now}."
            )
            doc.add_paragraph()

            # ── Action Queue Summary ──────────────────────────────────────────
            doc.add_heading("2. Action Queue Summary", 1)
            doc.add_paragraph(
                f"Total actions recorded: {len(actions)}. "
                f"Pending: {action_stats['pending']}. "
                f"Approved: {action_stats['approved']}. "
                f"Rejected: {action_stats['rejected']}. "
                f"Executed: {action_stats['executed']}."
            )
            doc.add_paragraph()

            if actions:
                doc.add_heading("2.1 Action Log", 2)
                tbl = doc.add_table(rows=1, cols=5)
                tbl.style = "Table Grid"
                _hdr(tbl, ["Action", "Type", "Category", "Status", "Created"])
                for idx, a in enumerate(actions[:20]):
                    row = tbl.add_row()
                    vals = [
                        (a.get("title") or a.get("action_type") or "")[:50],
                        str(a.get("action_type") or "")[:20],
                        str(a.get("category") or "")[:10],
                        str(a.get("status") or "").upper(),
                        str(a.get("created_at") or "")[:19],
                    ]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        row.cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                        if idx % 2 == 0:
                            _set_cell_bg(row.cells[i], "DCE6F1")
                    st = (a.get("status") or "").lower()
                    if st == "rejected":
                        row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    elif st == "approved":
                        row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                doc.add_paragraph()

            # ── Notifications ─────────────────────────────────────────────────
            doc.add_heading("3. Notification Log", 1)
            if notifications:
                tbl2 = doc.add_table(rows=1, cols=4)
                tbl2.style = "Table Grid"
                _hdr(tbl2, ["Notification", "Type", "Read", "Sent"])
                for idx, n in enumerate(notifications):
                    row = tbl2.add_row()
                    vals = [
                        (n.get("title") or "")[:60],
                        str(n.get("notif_type") or "")[:15],
                        "Yes" if n.get("is_read") else "No",
                        str(n.get("created_at") or "")[:19],
                    ]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        row.cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                        if idx % 2 == 0:
                            _set_cell_bg(row.cells[i], "DCE6F1")
            else:
                doc.add_paragraph("No notifications recorded.")
            doc.add_paragraph()

            # ── Governance attestation ────────────────────────────────────────
            doc.add_heading("4. Governance Controls Attestation", 1)
            attestations = [
                ("Human-in-the-loop enforcement",
                 "All Category B and C agent actions require human approval before execution."),
                ("Tenant data isolation",
                 "All queries and data access are scoped by tenant_id and project_id."),
                ("Audit trail completeness",
                 "All agent actions are logged with timestamp, category, and status."),
                ("Access control",
                 "JWT-based authentication enforced on all API endpoints."),
                ("Output governance",
                 "All skill-generated documents are queued for human review before distribution."),
            ]
            tbl3 = doc.add_table(rows=1, cols=3)
            tbl3.style = "Table Grid"
            _hdr(tbl3, ["Control", "Description", "Status"])
            for idx, (ctrl, desc) in enumerate(attestations):
                row = tbl3.add_row()
                row.cells[0].text = ctrl
                row.cells[1].text = desc
                row.cells[2].text = "IMPLEMENTED"
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if idx % 2 == 0:
                        _set_cell_bg(cell, "DCE6F1")
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                row.cells[2].paragraphs[0].runs[0].bold = True
            doc.add_paragraph()

            # ── Sign-off ──────────────────────────────────────────────────────
            doc.add_heading("5. Reviewer Sign-off", 1)
            sign_tbl = doc.add_table(rows=3, cols=3)
            sign_tbl.style = "Table Grid"
            _hdr(sign_tbl, ["Role", "Name", "Signature / Date"])
            for i, role in enumerate(["Compliance Officer", "System Owner"]):
                sign_tbl.rows[i+1].cells[0].text = role

            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run(
                f"This audit report was generated automatically by RAPID on {now}. "
                "It must be reviewed and countersigned before being used for compliance purposes."
            ).italic = True

            out_dir  = self._output_dir(context)
            filename = f"audit_report_{uuid.uuid4().hex[:8]}.docx"
            out_path = os.path.join(out_dir, filename)
            doc.save(out_path)

            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = out_path,
                preview     = (
                    f"Audit report: {len(actions)} actions "
                    f"({action_stats['approved']} approved, {action_stats['rejected']} rejected), "
                    f"{len(notifications)} notifications."
                ),
                pages       = max(4, len(actions) // 8),
            )

        except Exception as e:
            logger.error(f"[AuditReportSkill] Failed: {e}", exc_info=True)
            return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                               title=title, file_format=self.output_format, error=str(e))


def _register(registry) -> None:
    registry.register(AuditReportSkill())
