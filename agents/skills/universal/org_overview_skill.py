"""
Universal Skill: Organisation Overview Report (DOCX)

Aggregates all active projects across the tenant:
  - Project status summary table
  - Department health breakdown
  - Top risks across all projects
  - Headcount by department
  - Key metrics summary

Trigger phrases: "org overview", "organisation report", "organization overview",
                 "company overview", "portfolio overview", "all projects",
                 "tenant summary", "company report"
"""

from __future__ import annotations

import logging
import os
import sqlite3
import uuid
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.skills.base_skill import BaseSkill, SkillOutput
from infrastructure.project_context import ProjectContext

logger = logging.getLogger(__name__)


def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _header_row(table, headers: list[str], bg: str = "1F3564") -> None:
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[i], bg)


class OrgOverviewSkill(BaseSkill):
    skill_id        = "org_overview"
    dept_id         = "all"
    title_template  = "{project_name} — Organisation Overview"
    description     = "Universal: Generate a cross-tenant org overview DOCX — all projects, dept health, top risks, headcount."
    output_format   = "docx"
    trigger_phrases = [
        "org overview", "organisation report", "organization overview",
        "company overview", "portfolio overview", "all projects report",
        "tenant summary", "company report", "organisation summary",
        "cross-project report", "executive overview",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        title  = params.get("title") or self._make_title(context)

        try:
            import config

            # ── Collect all projects for this tenant ──────────────────────────
            # projects table has name/status/dept; project_registry has db_path
            reg_conn = sqlite3.connect(config.DB_PATH, timeout=10)
            reg_conn.row_factory = sqlite3.Row
            projects = [dict(r) for r in reg_conn.execute(
                """
                SELECT pr.project_id,
                       COALESCE(p.name, pr.project_id)       AS name,
                       COALESCE(p.status, pr.status)         AS status,
                       COALESCE(p.primary_dept_id, '')       AS dept_id,
                       pr.db_path,
                       COALESCE(p.created_at, pr.provisioned_at) AS created_at
                FROM project_registry pr
                LEFT JOIN projects p ON pr.project_id = p.project_id
                WHERE pr.tenant_id=? AND pr.status != 'archived'
                ORDER BY created_at DESC
                """,
                (context.tenant_id,),
            ).fetchall()]
            reg_conn.close()

            # ── Gather per-project data ───────────────────────────────────────
            all_risks:  list[dict] = []
            dept_kpis:  dict[str, list] = {}
            status_counts = {"active": 0, "at_risk": 0, "off_track": 0, "completed": 0, "other": 0}

            for proj in projects:
                st = (proj.get("status") or "other").lower().replace("-", "_")
                status_counts[st if st in status_counts else "other"] += 1

                db_path = proj.get("db_path") or ""
                if not db_path or not os.path.exists(db_path):
                    continue
                try:
                    pconn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=5)
                    pconn.row_factory = sqlite3.Row

                    risks = [dict(r) for r in pconn.execute(
                        "SELECT title, severity, status, owner FROM project_risks "
                        "WHERE severity IN ('high','critical') ORDER BY severity DESC LIMIT 3"
                    ).fetchall()]
                    for r in risks:
                        r["project_name"] = proj.get("name", proj["project_id"][:8])
                    all_risks.extend(risks)

                    dept = proj.get("dept_id") or "general"
                    kpis = [dict(r) for r in pconn.execute(
                        "SELECT name, current_value, target_value, unit, status FROM project_kpis LIMIT 3"
                    ).fetchall()]
                    if kpis:
                        dept_kpis.setdefault(dept, []).extend(kpis)

                    pconn.close()
                except Exception:
                    pass

            # ── People headcount ──────────────────────────────────────────────
            headcount: dict[str, int] = {}
            try:
                from infrastructure.people_directory import get_people_directory
                headcount = get_people_directory().dept_headcount(context.tenant_id)
            except Exception:
                pass

            # ── Build DOCX ────────────────────────────────────────────────────
            doc = Document()
            now = datetime.utcnow().strftime("%B %d, %Y")

            h = doc.add_heading(title, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Organisation Intelligence Report  |  Generated: {now}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()

            # ── Executive summary ─────────────────────────────────────────────
            doc.add_heading("Executive Summary", 1)
            total = len(projects)
            high_risk = len([r for r in all_risks if (r.get("severity") or "").lower() in ("high","critical")])
            doc.add_paragraph(
                f"Across {total} projects in this organisation, {status_counts.get('active',0)} are "
                f"active, {status_counts.get('at_risk',0)} are at risk, and "
                f"{status_counts.get('off_track',0)} are off-track. "
                f"There are {high_risk} high/critical risks requiring leadership attention."
            )
            doc.add_paragraph()

            # ── Project status table ──────────────────────────────────────────
            doc.add_heading("Project Portfolio", 1)
            if projects:
                tbl = doc.add_table(rows=1, cols=4)
                tbl.style = "Table Grid"
                _header_row(tbl, ["Project", "Department", "Status", "Created"])
                for idx, proj in enumerate(projects[:20]):
                    row = tbl.add_row()
                    vals = [
                        proj.get("name", proj["project_id"][:12]),
                        proj.get("dept_id") or "—",
                        (proj.get("status") or "").replace("_"," ").title(),
                        (proj.get("created_at") or "")[:10],
                    ]
                    for i, v in enumerate(vals):
                        row.cells[i].text = v
                        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                        if idx % 2 == 0:
                            _set_cell_bg(row.cells[i], "DCE6F1")
                    st = (proj.get("status") or "").lower()
                    if st in ("at_risk", "off_track"):
                        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            else:
                doc.add_paragraph("No projects found.")
            doc.add_paragraph()

            # ── Top risks ─────────────────────────────────────────────────────
            doc.add_heading("Top Risks Across Organisation", 1)
            if all_risks:
                tbl2 = doc.add_table(rows=1, cols=4)
                tbl2.style = "Table Grid"
                _header_row(tbl2, ["Risk", "Project", "Severity", "Status"])
                for idx, r in enumerate(all_risks[:15]):
                    row = tbl2.add_row()
                    vals = [r.get("title",""), r.get("project_name",""), r.get("severity",""), r.get("status","")]
                    for i, v in enumerate(vals):
                        row.cells[i].text = str(v)
                        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                        if idx % 2 == 0:
                            _set_cell_bg(row.cells[i], "DCE6F1")
                    if str(r.get("severity","")).lower() in ("high","critical"):
                        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        row.cells[2].paragraphs[0].runs[0].bold = True
            else:
                doc.add_paragraph("No high/critical risks identified.")
            doc.add_paragraph()

            # ── Headcount ─────────────────────────────────────────────────────
            if headcount:
                doc.add_heading("Department Headcount", 1)
                tbl3 = doc.add_table(rows=1, cols=2)
                tbl3.style = "Table Grid"
                _header_row(tbl3, ["Department", "Headcount"])
                for idx, (dept, cnt) in enumerate(sorted(headcount.items(), key=lambda x: -x[1])):
                    row = tbl3.add_row()
                    row.cells[0].text = dept
                    row.cells[1].text = str(cnt)
                    if idx % 2 == 0:
                        for cell in row.cells:
                            _set_cell_bg(cell, "DCE6F1")
                doc.add_paragraph()

            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Generated automatically by RAPID Organisation Intelligence.").italic = True

            out_dir  = self._output_dir(context)
            filename = f"org_overview_{uuid.uuid4().hex[:8]}.docx"
            out_path = os.path.join(out_dir, filename)
            doc.save(out_path)

            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = out_path,
                preview     = f"Org overview: {total} projects, {high_risk} critical risks.",
                pages       = max(3, total // 5),
            )

        except Exception as e:
            logger.error(f"[OrgOverviewSkill] Failed: {e}", exc_info=True)
            return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                               title=title, file_format=self.output_format, error=str(e))


def _register(registry) -> None:
    registry.register(OrgOverviewSkill())
