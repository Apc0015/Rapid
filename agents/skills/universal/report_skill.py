"""
Universal Skill: Report (DOCX)

Produces a professional Word document status report for any project.
Wraps document_engine.DocumentGenerator and extends it with:
  - action_plan report type
  - richer KPI + milestone + risk sections driven from live project data

Trigger phrases: "status report", "generate report", "project report",
                 "write a report", "action plan", "risk report"
"""

from __future__ import annotations

import logging
import os
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from agents.skills.base_skill import BaseSkill, SkillOutput
from infrastructure.project_context import ProjectContext

logger = logging.getLogger(__name__)


class ReportSkill(BaseSkill):
    skill_id        = "status_report"
    dept_id         = "all"
    title_template  = "{project_name} — Status Report"
    description     = "Generate a professional DOCX status report covering KPIs, milestones, budget, and risks."
    output_format   = "docx"
    trigger_phrases = [
        "status report", "generate report", "project report",
        "write a report", "produce a report", "create a report",
        "health report", "progress report",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        report_type = params.get("report_type", "status_report")
        title       = params.get("title") or self._make_title(context)

        try:
            from infrastructure.document_engine import get_document_generator
            gen    = get_document_generator(context)
            result = await gen.generate(
                report_type  = report_type,
                title        = title,
                produced_by  = "report_skill",
            )
            if not result.success:
                return SkillOutput(
                    skill_id    = self.skill_id,
                    dept_id     = self.dept_id,
                    title       = title,
                    file_format = self.output_format,
                    error       = result.error,
                )
            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = result.file_path,
                preview     = f"Report generated: {title} ({result.pages} pages)",
                pages       = result.pages,
            )
        except Exception as e:
            logger.error(f"[ReportSkill] Failed: {e}")
            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                error       = str(e),
            )


class RiskReportSkill(BaseSkill):
    skill_id        = "risk_report"
    dept_id         = "all"
    title_template  = "{project_name} — Risk Assessment Report"
    description     = "Generate a DOCX risk assessment with full risk register, impact analysis, and mitigation table."
    output_format   = "docx"
    trigger_phrases = [
        "risk report", "risk assessment", "risk register",
        "risks and mitigations", "generate risk report",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        title  = params.get("title") or self._make_title(context)
        try:
            from infrastructure.document_engine import get_document_generator
            gen    = get_document_generator(context)
            result = await gen.generate(
                report_type = "risk_report",
                title       = title,
                produced_by = "risk_report_skill",
            )
            if not result.success:
                return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                                   title=title, file_format=self.output_format, error=result.error)
            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = result.file_path,
                preview     = f"Risk report: {title} ({result.pages} pages)",
                pages       = result.pages,
            )
        except Exception as e:
            return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                               title=title, file_format=self.output_format, error=str(e))


class ActionPlanSkill(BaseSkill):
    skill_id        = "action_plan"
    dept_id         = "all"
    title_template  = "{project_name} — Action Plan"
    description     = "Generate a DOCX action plan listing what needs to happen, owners, and deadlines."
    output_format   = "docx"
    trigger_phrases = [
        "action plan", "generate action plan", "what needs to happen",
        "next steps document", "action items document",
    ]

    async def execute(self, context: ProjectContext, params: dict[str, Any] = None) -> SkillOutput:
        params = params or {}
        title  = params.get("title") or self._make_title(context)
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc  = Document()
            proj = context.project_id[:8]
            now  = datetime.utcnow().strftime("%B %d, %Y")

            # Title
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Generated: {now}  |  Project: {context.project_id}")
            doc.add_paragraph()

            # Fetch pending actions from DB
            db_path = getattr(context, "db_path", None)
            action_rows = []
            if db_path and os.path.exists(db_path):
                try:
                    conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=10)
                    conn.row_factory = sqlite3.Row
                    action_rows = conn.execute(
                        "SELECT title, description, priority, status, created_at FROM agent_action_queue "
                        "WHERE status IN ('pending','approved') ORDER BY priority DESC, created_at ASC LIMIT 30"
                    ).fetchall()
                    conn.close()
                except Exception:
                    pass

            # Fetch open milestones
            milestone_rows = []
            if db_path and os.path.exists(db_path):
                try:
                    conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=10)
                    conn.row_factory = sqlite3.Row
                    milestone_rows = conn.execute(
                        "SELECT name, due_date, owner, status FROM project_milestones "
                        "WHERE status != 'completed' ORDER BY due_date ASC LIMIT 20"
                    ).fetchall()
                    conn.close()
                except Exception:
                    pass

            # Pending actions section
            doc.add_heading("Pending Agent Actions", level=1)
            if action_rows:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for i, h in enumerate(["Action", "Priority", "Status", "Created"]):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].bold = True
                for r in action_rows:
                    row = table.add_row().cells
                    row[0].text = str(r["title"])
                    row[1].text = str(r["priority"])
                    row[2].text = str(r["status"])
                    row[3].text = str(r["created_at"])[:10]
            else:
                doc.add_paragraph("No pending actions at this time.")

            doc.add_paragraph()

            # Open milestones section
            doc.add_heading("Open Milestones", level=1)
            if milestone_rows:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for i, h in enumerate(["Milestone", "Due Date", "Owner", "Status"]):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].bold = True
                for r in milestone_rows:
                    row = table.add_row().cells
                    row[0].text = str(r["name"])
                    row[1].text = str(r["due_date"] or "TBD")
                    row[2].text = str(r["owner"] or "Unassigned")
                    row[3].text = str(r["status"])
            else:
                doc.add_paragraph("No open milestones.")

            # Save
            out_dir  = self._output_dir(context)
            filename = f"action_plan_{uuid.uuid4().hex[:8]}.docx"
            out_path = os.path.join(out_dir, filename)
            doc.save(out_path)

            return SkillOutput(
                skill_id    = self.skill_id,
                dept_id     = self.dept_id,
                title       = title,
                file_format = self.output_format,
                file_path   = out_path,
                preview     = f"Action plan with {len(action_rows)} actions and {len(milestone_rows)} milestones.",
                pages       = 2,
            )
        except Exception as e:
            logger.error(f"[ActionPlanSkill] Failed: {e}")
            return SkillOutput(skill_id=self.skill_id, dept_id=self.dept_id,
                               title=title, file_format=self.output_format, error=str(e))


# ── Self-registration ─────────────────────────────────────────────────────────

def _register(registry) -> None:
    registry.register(ReportSkill())
    registry.register(RiskReportSkill())
    registry.register(ActionPlanSkill())
